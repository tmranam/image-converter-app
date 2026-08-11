import streamlit as tf
import pandas as pd
from docx import Document
from io import BytesIO
import easyocr  # Alternative: import pytesseract
from PIL import Image

# Initialize the OCR reader (caches the model for performance)
@tf.cache_resource
def load_ocr():
    return easyocr.Reader(['en'])

reader = load_ocr()

# App Title & Layout
tf.set_page_config(page_title="Image Data Extractor", layout="centered")
tf.title("📷 Image Data Extractor & Converter")
tf.write("Upload an image to extract text or tabular data into Word or Excel formats.")

# 1. File Uploader Widget
uploaded_file = tf.file_uploader("Choose an image file...", type=["jpg", "jpeg", "png"])

if uploaded_file is not None:
    # Display the uploaded image
    image = Image.open(uploaded_file)
    tf.image(image, caption="Uploaded Image", use_container_width=True)
    
    tf.subheader("Conversion Options")
    tf.write("Select the formats you want to generate:")
    
    # 2. Tick Box (Checkbox) Widgets
    want_word = tf.checkbox("Convert to Word (.docx) - Best for paragraphs/text")
    want_excel = tf.checkbox("Convert to Excel (.xlsx) - Best for tables/grids")
    
    # 3. Action Button
    if tf.button("Extract and Convert"):
        if not want_word and not want_excel:
            tf.warning("Please select at least one format checkbox.")
        else:
            with tf.spinner("Processing image and extracting data..."):
                # Run OCR to extract text strings
                # EasyOCR returns a list of tuples: (bounding box, text, confidence)
                image_bytes = uploaded_file.getvalue()
                ocr_results = reader.readtext(image_bytes, detail=0) 
                
                # Combine extracted text lines
                extracted_text = "\n".join(ocr_results)
                
                if not extracted_text.strip():
                    tf.error("No text could be detected in the image.")
                else:
                    tf.success("Data successfully extracted!")
                    
                    # 4. Generate Requested Formats
                    if want_word:
                        # Create Word document in memory
                        doc = Document()
                        doc.add_heading("Extracted Text Document", level=1)
                        doc.add_paragraph(extracted_text)
                        
                        word_buffer = BytesIO()
                        doc.save(word_buffer)
                        word_buffer.seek(0)
                        
                        # Streamlit Download Button for Word
                        tf.download_button(
                            label="📥 Download Word Document",
                            data=word_buffer,
                            file_name="extracted_data.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
                        
                    if want_excel:
                        # For a simple Excel conversion, split rows by newline
                        # For advanced tables, consider splitting rows by commas/spaces
                        rows = [line.split() for line in ocr_results if line.strip()]
                        df = pd.DataFrame(rows)
                        
                        excel_buffer = BytesIO()
                        with pd.ExcelWriter(excel_buffer, engine='xlsxwriter') as writer_engine:
                            df.to_excel(writer_engine, index=False, header=False, sheet_name="Extracted Data")
                        excel_buffer.seek(0)
                        
                        # Streamlit Download Button for Excel
                        tf.download_button(
                            label="📥 Download Excel Spreadsheet",
                            data=excel_buffer,
                            file_name="extracted_data.xlsx",
                            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                        )
